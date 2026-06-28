import os
import uuid
import logging
from datetime import datetime
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.core.deps import get_current_user
from app.domains.knowledge import models, schemas
from app.domains.auth.models import User

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/knowledge", tags=["Knowledge"])

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "../../../static/uploads/knowledge")
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {
    ".pdf", ".txt", ".text", ".docx", ".doc",
    ".xlsx", ".xls", ".csv", ".json", ".md",
}

EXT_TO_TYPE = {
    ".pdf":   models.DocumentType.PDF,
    ".docx":  models.DocumentType.WORD,
    ".doc":   models.DocumentType.WORD,
    ".xlsx":  models.DocumentType.EXCEL,
    ".xls":   models.DocumentType.EXCEL,
    ".csv":   models.DocumentType.CSV,
    ".txt":   models.DocumentType.TEXT,
    ".text":  models.DocumentType.TEXT,
    ".json":  models.DocumentType.JSON,
    ".md":    models.DocumentType.TEXT,
}


def _extract_text(file_path: str, ext: str) -> str:
    try:
        if ext == ".pdf":
            import pdfplumber
            texts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        texts.append(t)
            return "\n\n".join(texts)

        elif ext in (".docx", ".doc"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"=== {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    row_str = "\t".join(str(v) if v is not None else "" for v in row)
                    if row_str.strip():
                        lines.append(row_str)
            return "\n".join(lines)

        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        elif ext == ".json":
            import json
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)
            return json.dumps(data, ensure_ascii=False, indent=2)

        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

    except Exception as e:
        logger.warning(f"Text extraction failed for {file_path}: {e}")
        return ""


# ─── CATEGORIES ───────────────────────────────────────────────────────────────

@router.get("/categories", response_model=list[schemas.CategoryOut])
def list_categories(db: Session = Depends(get_db), _: User = Depends(get_current_user)):
    return db.query(models.KnowledgeCategory).all()


@router.post("/categories", response_model=schemas.CategoryOut)
def create_category(
    data: schemas.CategoryCreate,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    cat = models.KnowledgeCategory(**data.model_dump())
    db.add(cat)
    db.commit()
    db.refresh(cat)
    return cat


@router.put("/categories/{cat_id}", response_model=schemas.CategoryOut)
def update_category(
    cat_id: int,
    data: schemas.CategoryUpdate,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    cat = db.query(models.KnowledgeCategory).filter(models.KnowledgeCategory.id == cat_id).first()
    if not cat:
        raise HTTPException(status_code=404, detail="Category not found")
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(cat, field, value)
    db.commit()
    db.refresh(cat)
    return cat


@router.delete("/categories/{cat_id}")
def delete_category(
    cat_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    cat = db.query(models.KnowledgeCategory).filter(models.KnowledgeCategory.id == cat_id).first()
    if not cat:
        raise HTTPException(status_code=404, detail="Category not found")
    db.delete(cat)
    db.commit()
    return {"message": "Category deleted"}


# ─── DOCUMENTS ────────────────────────────────────────────────────────────────

@router.get("/documents", response_model=list[schemas.DocumentOut])
def list_documents(
    skip: int = 0,
    limit: int = 100,
    search: str | None = Query(None),
    category_id: int | None = Query(None),
    status: str | None = Query(None),
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    q = db.query(models.KnowledgeDocument)
    if search:
        q = q.filter(
            models.KnowledgeDocument.title.ilike(f"%{search}%") |
            models.KnowledgeDocument.title_ar.ilike(f"%{search}%") |
            models.KnowledgeDocument.content.ilike(f"%{search}%")
        )
    if category_id:
        q = q.filter(models.KnowledgeDocument.category_id == category_id)
    if status:
        q = q.filter(models.KnowledgeDocument.status == status)
    return q.order_by(models.KnowledgeDocument.created_at.desc()).offset(skip).limit(limit).all()


@router.post("/documents", response_model=schemas.DocumentOut)
def create_document(
    data: schemas.DocumentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = models.KnowledgeDocument(**data.model_dump(), created_by=current_user.id)
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


@router.post("/documents/upload", response_model=schemas.DocumentOut)
async def upload_document(
    title_ar: str = Form(...),
    title: str = Form(None),
    category_id: int = Form(None),
    summary: str = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    original_name = file.filename or "document"
    ext = os.path.splitext(original_name)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"نوع الملف غير مدعوم. الأنواع المدعومة: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    unique_name = f"{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(UPLOAD_DIR, unique_name)

    content_bytes = await file.read()
    file_size = len(content_bytes)
    with open(save_path, "wb") as f:
        f.write(content_bytes)

    extracted = _extract_text(save_path, ext)

    doc_type = EXT_TO_TYPE.get(ext, models.DocumentType.OTHER)

    doc = models.KnowledgeDocument(
        title=title or title_ar,
        title_ar=title_ar,
        content=extracted or None,
        summary=summary or None,
        doc_type=doc_type,
        status=models.KnowledgeStatus.ACTIVE,
        file_path=save_path,
        file_name=original_name,
        file_size=file_size,
        version="1.0",
        category_id=category_id or None,
        created_by=current_user.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


@router.get("/documents/{doc_id}", response_model=schemas.DocumentOut)
def get_document(
    doc_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    doc = db.query(models.KnowledgeDocument).filter(models.KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.put("/documents/{doc_id}", response_model=schemas.DocumentOut)
def update_document(
    doc_id: int,
    data: schemas.DocumentUpdate,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    doc = db.query(models.KnowledgeDocument).filter(models.KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(doc, field, value)
    doc.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(doc)
    return doc


@router.post("/documents/{doc_id}/train")
def train_document(
    doc_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    """
    Full RAG pipeline v3.0:
    Clean → Section Detect → Semantic Chunk (with overlap) →
    Per-chunk Keywords + Questions → Entities → Facts → Intent →
    Save chunks to knowledge_chunks table for BM25 chunk-level retrieval.
    """
    from app.domains.knowledge.pipeline import build_training_json, build_chunks_for_db

    doc = db.query(models.KnowledgeDocument).filter(models.KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status != models.KnowledgeStatus.ACTIVE:
        raise HTTPException(status_code=400, detail="يجب أن تكون الوثيقة في حالة نشطة للتدريب")

    # Run full pipeline
    training_json = build_training_json(doc)

    # Delete old chunks (re-train)
    db.query(models.KnowledgeChunk).filter(
        models.KnowledgeChunk.document_id == doc_id
    ).delete(synchronize_session=False)

    # Save new chunks to DB for fast retrieval
    chunks_data = build_chunks_for_db(doc, training_json)
    for cd in chunks_data:
        chunk_obj = models.KnowledgeChunk(**cd)
        db.add(chunk_obj)

    # Update document
    doc.content = training_json
    doc.is_trained = True
    doc.trained_at = datetime.utcnow()
    doc.updated_at = datetime.utcnow()
    db.commit()

    import json as _json
    try:
        parsed = _json.loads(training_json)
        stats = parsed.get("stats", {})
        pipeline_info = parsed.get("pipeline", {})
    except Exception:
        stats = {}
        pipeline_info = {}

    return {
        "message": "تم تدريب الوكيل على هذه الوثيقة بنجاح (Pipeline v3.0)",
        "doc_id": doc_id,
        "pipeline": {
            "version":       pipeline_info.get("version", "3.0"),
            "chunks":        stats.get("chunk_count", 0),
            "keywords":      stats.get("keyword_count", 0),
            "facts":         stats.get("fact_count", 0),
            "questions":     stats.get("question_count", 0),
            "intent":        parsed.get("intent") if training_json else None,
            "language":      parsed.get("language") if training_json else None,
            "stages":        pipeline_info.get("stages", []),
        },
        "chunks_saved_to_db": len(chunks_data),
    }


@router.post("/documents/{doc_id}/untrain")
def untrain_document(
    doc_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    doc = db.query(models.KnowledgeDocument).filter(models.KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete stored chunks
    db.query(models.KnowledgeChunk).filter(
        models.KnowledgeChunk.document_id == doc_id
    ).delete(synchronize_session=False)

    doc.is_trained = False
    doc.trained_at = None
    doc.updated_at = datetime.utcnow()
    db.commit()
    return {"message": "تم إلغاء تدريب الوكيل على هذه الوثيقة", "doc_id": doc_id}


@router.delete("/documents/{doc_id}")
def delete_document(
    doc_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    doc = db.query(models.KnowledgeDocument).filter(models.KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except Exception:
            pass
    db.delete(doc)
    db.commit()
    return {"message": "تم حذف الوثيقة"}


# ─── CHUNKS ───────────────────────────────────────────────────────────────────

@router.get("/documents/{doc_id}/chunks")
def get_document_chunks(
    doc_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    """Inspect all chunks and their generated questions for a trained document."""
    doc = db.query(models.KnowledgeDocument).filter(models.KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    chunks = (
        db.query(models.KnowledgeChunk)
        .filter(models.KnowledgeChunk.document_id == doc_id)
        .order_by(models.KnowledgeChunk.chunk_index)
        .all()
    )
    return {
        "doc_id":      doc_id,
        "title":       doc.title_ar or doc.title,
        "chunk_count": len(chunks),
        "chunks": [
            {
                "id":              ch.id,
                "index":           ch.chunk_index,
                "text":            ch.text,
                "section":         ch.section_heading,
                "keywords":        [k["term"] for k in (ch.keywords or [])][:8],
                "questions":       ch.questions or [],
                "word_count":      ch.word_count,
                "importance":      ch.importance_score,
                "retrieval_count": ch.retrieval_count,
            }
            for ch in chunks
        ],
    }


# ─── RAG SEARCH ───────────────────────────────────────────────────────────────

@router.get("/rag/search")
def rag_search(
    q: str = Query(..., min_length=2),
    top_k: int = Query(5, ge=1, le=20),
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    """
    Hybrid BM25 + question-match search across trained documents.
    v3.0: searches at chunk level, returns confidence scores.
    """
    from app.domains.knowledge.pipeline import rag_search as _rag_search

    trained_docs = (
        db.query(models.KnowledgeDocument)
        .filter(models.KnowledgeDocument.is_trained == True)
        .all()
    )
    results = _rag_search(q, trained_docs, top_k=top_k, db=db)

    # Update retrieval stats for matched docs
    if results:
        matched_ids = [r["doc_id"] for r in results]
        now = datetime.utcnow()
        db.query(models.KnowledgeDocument).filter(
            models.KnowledgeDocument.id.in_(matched_ids)
        ).update(
            {
                models.KnowledgeDocument.retrieval_count: models.KnowledgeDocument.retrieval_count + 1,
                models.KnowledgeDocument.last_retrieved_at: now,
            },
            synchronize_session=False,
        )
        # Also increment chunk retrieval counts
        for r in results:
            if r.get("matched_questions"):
                pass  # chunk_id tracking can be added later
        db.commit()

    return {
        "query":    q,
        "expanded": True,
        "results":  results,
        "total":    len(results),
    }


# ─── FEEDBACK LOOP ────────────────────────────────────────────────────────────

@router.post("/rag/feedback")
def submit_feedback(
    query: str,
    doc_id: int | None = None,
    chunk_id: int | None = None,
    was_helpful: bool | None = None,
    feedback_text: str | None = None,
    confidence_shown: str | None = None,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    """
    Record whether a retrieved result was helpful.
    Feedback is used to track quality and can be used to boost chunk importance.
    """
    fb = models.KnowledgeFeedback(
        query=query,
        doc_id=doc_id,
        chunk_id=chunk_id,
        was_helpful=was_helpful,
        feedback_text=feedback_text,
        confidence_shown=confidence_shown,
    )
    db.add(fb)

    # Boost chunk importance if helpful
    if was_helpful and chunk_id:
        chunk = db.query(models.KnowledgeChunk).filter(
            models.KnowledgeChunk.id == chunk_id
        ).first()
        if chunk:
            chunk.importance_score = min(chunk.importance_score + 0.05, 2.0)
            chunk.retrieval_count += 1
    elif was_helpful is False and chunk_id:
        chunk = db.query(models.KnowledgeChunk).filter(
            models.KnowledgeChunk.id == chunk_id
        ).first()
        if chunk:
            chunk.importance_score = max(chunk.importance_score - 0.03, 0.5)

    db.commit()
    return {"message": "تم تسجيل التقييم بنجاح", "feedback_id": fb.id}


# ─── PIPELINE SUMMARY ─────────────────────────────────────────────────────────

@router.get("/pipeline/summary")
def pipeline_summary(
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    """Overview of RAG pipeline v3.0 processing across all trained documents."""
    import json as _json

    trained = (
        db.query(models.KnowledgeDocument)
        .filter(models.KnowledgeDocument.is_trained == True)
        .order_by(models.KnowledgeDocument.trained_at.desc())
        .all()
    )

    docs_info = []
    total_chunks = 0
    total_facts = 0
    total_keywords = 0
    total_questions = 0
    intents: dict = {}
    languages: dict = {}

    for doc in trained:
        # Count DB chunks for this doc
        db_chunk_count = db.query(models.KnowledgeChunk).filter(
            models.KnowledgeChunk.document_id == doc.id
        ).count()

        info: dict = {
            "id":              doc.id,
            "title":           doc.title_ar or doc.title,
            "type":            str(doc.doc_type.value) if doc.doc_type else "other",
            "trained_at":      doc.trained_at.isoformat() if doc.trained_at else None,
            "pipeline_v":      "1.0",
            "intent":          None,
            "language":        None,
            "chunks":          db_chunk_count,
            "facts":           0,
            "keywords":        0,
            "questions":       0,
            "retrieval_count": doc.retrieval_count or 0,
            "stages":          [],
        }

        if doc.content:
            try:
                data = _json.loads(doc.content)
                stats = data.get("stats", {})
                info["intent"]      = data.get("intent")
                info["language"]    = data.get("language")
                info["facts"]       = stats.get("fact_count", 0)
                info["keywords"]    = stats.get("keyword_count", 0)
                info["questions"]   = stats.get("question_count", 0)
                info["pipeline_v"]  = data.get("pipeline", {}).get("version", "1.0")
                info["stages"]      = data.get("pipeline", {}).get("stages", [])
                total_facts       += info["facts"]
                total_keywords    += info["keywords"]
                total_questions   += info["questions"]
                if info["intent"]:
                    intents[info["intent"]] = intents.get(info["intent"], 0) + 1
                if info["language"]:
                    languages[info["language"]] = languages.get(info["language"], 0) + 1
            except Exception:
                pass

        total_chunks += db_chunk_count
        docs_info.append(info)

    return {
        "overview": {
            "total_trained":      len(trained),
            "total_chunks":       total_chunks,
            "total_facts":        total_facts,
            "total_keywords":     total_keywords,
            "total_questions":    total_questions,
            "intent_breakdown":   intents,
            "language_breakdown": languages,
            "pipeline_version":   "3.0",
        },
        "documents": docs_info,
    }
